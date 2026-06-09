"""
korodur_referenz — Renderer für KORODUR One-Page-Referenzen.

Layout (A4 hochkant, ein Blatt):
  - Header-Strip: Mini-Logo rechts (~22 mm), Headline links auf Logo-Höhe
  - Schmaler Divider (Navy 0.5pt)
  - Hero-Bild links (~60 % Breite, 4:3) + Sidebar-Karte rechts mit Navy-Stripe
    und 5 Schlaglichtern (Auftrag, Fläche, Auftragsstärke, Bauzeit, Beobachtung)
  - Lead-Absatz (11 pt, Navy)
  - Story-Absätze (10 pt)
  - Bildstreifen: 3 Bilder + kurze Bildunterschrift (3:2)
  - Word-Footer: einzeilige Firmen-/Kontaktzeile links, QR-Code rechts

CD-Werte:
  - Schrift: Arial
  - Navy: #002D59 (Titel, Akzent)
  - Dunkelblau: #1E3A5F (Footer-Text)
  - Adresse: Wernher-von-Braun Str. 4, 92224 Amberg
  - Tel: +49 (0) 96 21 - 47 59 0
  - Web: www.korodur.de
"""
from __future__ import annotations

import hashlib
import io
import os
import tempfile
from typing import Sequence, Tuple

import qrcode
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from PIL import Image, ImageOps


_ORIENT_CACHE: dict[str, str] = {}


def _orient(path: str) -> str:
    """Wendet EXIF-Orientation an und gibt einen Pfad zu einem korrekt
    ausgerichteten JPG zurück. Notwendig, weil python-docx EXIF-Tags
    ignoriert und Bilder sonst um 90° gedreht eingebettet werden.

    Cached pro Originalpfad, damit derselbe Input nur einmal konvertiert wird.
    """
    abs_path = os.path.abspath(path)
    if abs_path in _ORIENT_CACHE:
        cached = _ORIENT_CACHE[abs_path]
        if os.path.exists(cached):
            return cached
    im = Image.open(abs_path)
    im = ImageOps.exif_transpose(im)
    if im.mode not in ("RGB", "L"):
        im = im.convert("RGB")
    # Resize auf max. 1600 px Längsseite, damit die DOCX nicht aufgebläht wird
    max_side = 1600
    if max(im.size) > max_side:
        im.thumbnail((max_side, max_side), Image.LANCZOS)
    h = hashlib.sha1(abs_path.encode()).hexdigest()[:10]
    tmp_dir = tempfile.gettempdir()
    out = os.path.join(tmp_dir, f"korodur_oriented_{h}.jpg")
    im.save(out, "JPEG", quality=88)
    _ORIENT_CACHE[abs_path] = out
    return out


# --- KORODUR Corporate Design ---
NAVY = RGBColor(0x00, 0x2D, 0x59)        # Akzent/Titel, kräftigstes Blau
DARKBLUE = RGBColor(0x1E, 0x3A, 0x5F)    # Body-Text und Footer (durchgängig)
GREY = RGBColor(0x8A, 0x93, 0xA6)        # Sidebar-Labels, Captions
TEXT = DARKBLUE                          # Backwards-Compat: TEXT = DARKBLUE,
                                         # damit Body-Text durchgängig dunkelblau
                                         # läuft (kein Schwarz/Anthrazit mehr).
NAVY_HEX = "002D59"
DARKBLUE_HEX = "1E3A5F"
GREY_HEX = "8A93A6"
SIDEBAR_FILL = "F4F6FB"                  # weicher hellblauer Hintergrund
DIVIDER_HEX = "D0D5E0"

DEFAULT_COMPANY_LINE = (
    "KORODUR International GmbH · Wernher-von-Braun-Str. 4 · "
    "92224 Amberg · +49 9621 4759-0 · korodur.de"
)
DEFAULT_SLUG_BASE = "https://korodur.de/r/"


def _asset(name: str) -> str:
    here = os.path.dirname(os.path.abspath(__file__))
    return os.path.normpath(os.path.join(here, "..", "assets", name))


# --- low-level helpers ----------------------------------------------------

def _set_arial(run):
    rPr = run._element.get_or_add_rPr()
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rFonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), "Arial")
    rPr.append(rFonts)


def _styled_run(p, text, *, bold=False, italic=False, size=None, color=None):
    run = p.add_run(text)
    _set_arial(run)
    if size is not None:
        run.font.size = Pt(size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color is not None:
        run.font.color.rgb = color
    return run


def _shade(cell, color_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _no_borders(table):
    """Komplett unsichtbare Tabellenränder."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _cant_split_row(row):
    """Verhindert, dass eine Tabellenreihe über zwei Seiten geteilt wird
    (Bild und Caption bleiben so immer zusammen)."""
    trPr = row._tr.get_or_add_trPr()
    existing = trPr.find(qn("w:cantSplit"))
    if existing is None:
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)


def _cell_left_border(cell, color_hex: str, size_eighths: int = 24):
    """Setzt einen Navy-Stripe links an die Sidebar-Zelle.
    size_eighths sind Achtel-Punkte (24 = 3 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_eighths))
    left.set(qn("w:color"), color_hex)
    borders.append(left)
    tcPr.append(borders)


def _cell_padding(cell, top=4, bottom=4, left=4, right=4):
    """Innenabstände einer Zelle in 'twips' (1/20 pt). 80 ≈ 4 pt."""
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcMar"))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                       ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(int(val * 20)))  # in twips
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_col_widths(table, widths_cm: Sequence[float]):
    """Tabellenspaltenbreiten fix setzen (Word ignoriert sonst gerne Cm())."""
    table.autofit = False
    for i, w in enumerate(widths_cm):
        for r in range(len(table.rows)):
            table.cell(r, i).width = Cm(w)
    # Auch tblGrid setzen, damit Word sich daran hält
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 567)))  # 1 cm = 567 twips
        tblGrid.append(gc)
    tbl.insert(0, tblGrid)


# --- Hauptklasse ----------------------------------------------------------


class KorodurRef:
    """One-Page-Referenz im KORODUR-Design.

    Beispiel:
        ref = KorodurRef()
        ref.header(kicker="Referenz · Logistikzentrum",
                   headline="Sanierte LKW-Zufahrt – nach 9 Monaten rissfrei")
        ref.divider()
        ref.hero_with_sidebar(
            hero_path="hero.jpg",
            facts=[
                ("Auftrag", "NBM-Hartstoffschicht"),
                ("Fläche", "1.250 m²"),
                ("Auftragsstärke", "8 mm"),
                ("Bauzeit", "2 Tage"),
                ("Beobachtung", "9 Monate rissfrei"),
            ],
        )
        ref.lead("In nur zwei Tagen wurde …")
        ref.paragraph("Die Aufgabe …")
        ref.paragraph("Die Lösung …")
        ref.image_strip([
            ("vorher.jpg", "Vorher · 10/2024"),
            ("auftrag.jpg", "NBM-Auftrag · 06/2025"),
            ("nachher.jpg", "Nachher · 03/2026"),
        ])
        ref.footer(slug="brummer-lkw-zufahrt")
        ref.save("teaser.docx")
    """

    PAGE_W = 21.0   # cm, A4
    PAGE_H = 29.7
    MARGIN_TOP = 1.4
    MARGIN_BOTTOM = 1.0
    MARGIN_LEFT = 1.8
    MARGIN_RIGHT = 1.8

    @property
    def CONTENT_W(self) -> float:
        return self.PAGE_W - self.MARGIN_LEFT - self.MARGIN_RIGHT  # 17.4 cm

    def __init__(self):
        self._doc = Document()
        self._setup_page()
        self._setup_default_paragraph_style()

    # --- Setup --------------------------------------------------------

    def _setup_page(self):
        section = self._doc.sections[0]
        section.page_width = Cm(self.PAGE_W)
        section.page_height = Cm(self.PAGE_H)
        section.top_margin = Cm(self.MARGIN_TOP)
        section.bottom_margin = Cm(self.MARGIN_BOTTOM)
        section.left_margin = Cm(self.MARGIN_LEFT)
        section.right_margin = Cm(self.MARGIN_RIGHT)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.5)

    def _setup_default_paragraph_style(self):
        # Kleinere Default-Abstände, damit auf einer Seite Platz ist
        styles = self._doc.styles
        normal = styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10)
        normal.paragraph_format.space_after = Pt(2)
        normal.paragraph_format.space_before = Pt(0)

    # --- Bausteine ----------------------------------------------------

    def header(self, *, kicker: str | None = None, headline: str,
               eyebrow: str | None = None, subline: str | None = None):
        """Header-Strip: Eyebrow + Headline links, Mini-Logo rechts auf gleicher Höhe.

        Args:
            eyebrow: Versalien-Tag oberhalb der Headline (z. B. "Referenz · Logistikzentrum").
                Alias für den älteren `kicker`-Parameter, beide werden unterstützt.
            headline: Die Hauptüberschrift, max. 2 Zeilen.
            subline: Optionale Pointe direkt unter der Headline. **Nur einsetzen,
                wenn die Headline ohne sie zu nackt wirkt.** Wenn ausgelassen,
                stehen Eyebrow + Headline allein.
            kicker: Alias zu `eyebrow` (Rückwärtskompatibilität).
        """
        # Eyebrow oder kicker akzeptieren
        eyebrow_text = eyebrow if eyebrow is not None else kicker

        # Rechte Spalte knapp am Logo (Mm(28) ≈ 2.8 cm) halten,
        # damit die Headline links maximal Platz hat.
        table = self._doc.add_table(rows=1, cols=2)
        _no_borders(table)
        _set_col_widths(table, [self.CONTENT_W - 3.2, 3.2])

        # Linke Zelle: Eyebrow + Headline (+ optionale Subline)
        c_left = table.cell(0, 0)
        c_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_padding(c_left, top=0, bottom=0, left=0, right=4)
        c_left.text = ""
        if eyebrow_text:
            p_k = c_left.paragraphs[0]
            p_k.paragraph_format.space_after = Pt(2)
            _styled_run(p_k, eyebrow_text.upper(), size=8, color=GREY)
            # Letter-spacing für Eyebrow (per OOXML w:spacing)
            for run in p_k.runs:
                rPr = run._element.get_or_add_rPr()
                spacing = OxmlElement("w:spacing")
                spacing.set(qn("w:val"), "30")  # in twentieths of a point
                rPr.append(spacing)
            p_h = c_left.add_paragraph()
        else:
            p_h = c_left.paragraphs[0]
        # Wenn Subline folgt, kleinen Abstand lassen
        p_h.paragraph_format.space_after = Pt(2 if subline else 0)
        # Etwas dezenter (15 pt statt 16) hält längere Headlines auf einer Zeile.
        # Manuelle Zeilenumbrüche via "\n" werden als harte Umbrüche gesetzt,
        # damit eine zu lange Headline gezielt sauber umbrechen kann.
        for idx, line in enumerate(headline.split("\n")):
            if idx > 0:
                br = p_h.add_run()
                _set_arial(br)
                br.add_break()
            _styled_run(p_h, line, bold=True, size=15, color=NAVY)

        if subline:
            p_s = c_left.add_paragraph()
            p_s.paragraph_format.space_after = Pt(0)
            _styled_run(p_s, subline, bold=True, size=11, color=NAVY)

        # Rechte Zelle: Logo
        # Logo etwas kleiner (Mm(25)) + kleines right-Padding, damit es
        # nicht am Seitenrand abgeschnitten wird.
        c_right = table.cell(0, 1)
        c_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_padding(c_right, top=0, bottom=0, left=4, right=2)
        p_logo = c_right.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_logo.paragraph_format.space_after = Pt(0)
        run = p_logo.add_run()
        _set_arial(run)
        run.add_picture(_asset("logo.png"), width=Mm(25))
        return table

    def divider(self):
        """Schmaler Trenner unterhalb des Headers."""
        p = self._doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")           # 0.5 pt
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), DIVIDER_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Sidebar-Labels je Sprache (DE Default, EN für internationale Varianten).
    _SIDEBAR_LABELS = {
        "de": ("PRODUKT", "AUF EINEN BLICK"),
        "en": ("PRODUCT", "QUICK OVERVIEW"),
    }

    def hero_with_sidebar(self, *,
                           facts: Sequence[Tuple[str, str]],
                           hero_path: str | None = None,
                           hero_pair: Tuple[str, str] | None = None,
                           product: str | None = None,
                           product_claim: str | None = None,
                           hero_captions: Tuple[str, str] | None = None,
                           lang: str = "de",
                           product_label: str | None = None,
                           section_label: str | None = None):
        """Hero-Bereich links, Sidebar-Karte rechts, gleiche Zeile.

        Entweder `hero_path` (ein Bild, volle Hero-Breite) ODER `hero_pair`
        (zwei Bilder Vorher/Nachher nebeneinander, je 50 %) übergeben.

        Wenn `product` gesetzt ist, wird oben in der Sidebar ein hervor­
        gehobener Produkt-Block (Pflichtelement jeder Referenz) gerendert:
            Produkt: <name>          (klein grau Label)
            <product>                (bold, größer, navy)
            <product_claim>          (1-Zeiler kursiv, dunkelblau, optional)
        Danach folgen ein dünner Trenner und die übrigen Schlaglichter.
        """
        if (hero_path is None) == (hero_pair is None):
            raise ValueError("hero_with_sidebar() braucht genau eines: hero_path ODER hero_pair.")

        # Sidebar-Label-Texte sprachabhängig wählen (explizite Overrides gewinnen).
        _def_prod, _def_sect = self._SIDEBAR_LABELS.get(lang, self._SIDEBAR_LABELS["de"])
        prod_label_text = product_label if product_label is not None else _def_prod
        sect_label_text = section_label if section_label is not None else _def_sect

        hero_w = round(self.CONTENT_W * 0.60, 2)     # ~10.4 cm
        side_w = round(self.CONTENT_W * 0.38, 2)     # ~6.6 cm
        gap = round(self.CONTENT_W - hero_w - side_w, 2)  # Rest als visuelle Lücke

        table = self._doc.add_table(rows=1, cols=3)
        _no_borders(table)
        _set_col_widths(table, [hero_w, gap, side_w])

        # Hero
        c_hero = table.cell(0, 0)
        c_hero.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _cell_padding(c_hero, top=0, bottom=0, left=0, right=0)
        p_hero = c_hero.paragraphs[0]
        p_hero.paragraph_format.space_after = Pt(0)

        if hero_path:
            run = p_hero.add_run()
            _set_arial(run)
            run.add_picture(_orient(hero_path), width=Cm(hero_w))
        else:
            # Zwei Bilder Vorher/Nachher nebeneinander
            inner_gap = 0.1
            half_w = round((hero_w - inner_gap) / 2, 2)
            pair_table = c_hero.add_table(rows=1, cols=3)
            _no_borders(pair_table)
            _set_col_widths(pair_table, [half_w, inner_gap, half_w])
            for i, path in enumerate(hero_pair):
                col = i * 2
                cell = pair_table.cell(0, col)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                _cell_padding(cell, top=0, bottom=0, left=0, right=0)
                p_img = cell.paragraphs[0]
                p_img.paragraph_format.space_after = Pt(0)
                r = p_img.add_run()
                _set_arial(r)
                r.add_picture(_orient(path), width=Cm(half_w))
            # Den ursprünglichen leeren Paragraph in c_hero entfernen,
            # damit kein doppelter Abstand entsteht
            p_hero._element.getparent().remove(p_hero._element)

        # Optional: Vorher/Nachher-Beschriftung unter dem Hero (50/50)
        if hero_captions:
            # Caption-Spalten exakt auf die beiden Bildhälften legen (kein 0.1-Gap),
            # damit die zentrierten Captions sauber unter ihrer Bildhälfte sitzen.
            half_w = round(hero_w / 2, 2)
            cap_table = c_hero.add_table(rows=1, cols=2)
            _no_borders(cap_table)
            _set_col_widths(cap_table, [half_w, half_w])
            for i, text in enumerate(hero_captions):
                cell = cap_table.cell(0, i)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                _cell_padding(cell, top=4, bottom=0, left=0, right=0)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(3)
                # Hero-Captions mittig unter ihrer Bildhälfte ausrichten,
                # damit sie bündig zu den (ebenfalls zentrierten) Streifen-Captions sind.
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _styled_run(p, text, size=8, color=GREY)

        # Spacer-Spalte
        c_gap = table.cell(0, 1)
        _cell_padding(c_gap, top=0, bottom=0, left=0, right=0)

        # Sidebar-Karte
        c_side = table.cell(0, 2)
        c_side.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _shade(c_side, SIDEBAR_FILL)
        _cell_left_border(c_side, NAVY_HEX, size_eighths=24)
        _cell_padding(c_side, top=8, bottom=8, left=10, right=8)

        first_paragraph_used = False

        def _next_paragraph():
            nonlocal first_paragraph_used
            if not first_paragraph_used:
                first_paragraph_used = True
                return c_side.paragraphs[0]
            return c_side.add_paragraph()

        # Optionaler Produkt-Block (Pflicht für Referenzen)
        if product:
            p_label = _next_paragraph()
            p_label.paragraph_format.space_after = Pt(0)
            _styled_run(p_label, prod_label_text, size=8, color=GREY)
            for r in p_label.runs:
                rPr = r._element.get_or_add_rPr()
                spacing = OxmlElement("w:spacing")
                spacing.set(qn("w:val"), "30")
                rPr.append(spacing)
            p_prod = _next_paragraph()
            p_prod.paragraph_format.space_after = Pt(0)
            _styled_run(p_prod, product, bold=True, size=12, color=NAVY)
            if product_claim:
                p_claim = _next_paragraph()
                p_claim.paragraph_format.space_after = Pt(8)
                _styled_run(p_claim, product_claim, italic=True, size=9,
                             color=DARKBLUE)
            else:
                _next_paragraph().paragraph_format.space_after = Pt(8)
            # dünner Trenner
            p_div = _next_paragraph()
            p_div.paragraph_format.space_before = Pt(0)
            p_div.paragraph_format.space_after = Pt(4)
            pPr = p_div._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "0")
            bottom.set(qn("w:color"), DIVIDER_HEX)
            pBdr.append(bottom)
            pPr.append(pBdr)

        # Sektions-Label "AUF EINEN BLICK"
        p_section = _next_paragraph()
        p_section.paragraph_format.space_after = Pt(4)
        p_section.paragraph_format.space_before = Pt(0)
        _styled_run(p_section, sect_label_text, size=8, color=GREY)
        for r in p_section.runs:
            rPr = r._element.get_or_add_rPr()
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:val"), "30")
            rPr.append(spacing)

        # Schlaglichter
        for label, value in facts:
            p_l = _next_paragraph()
            p_l.paragraph_format.space_after = Pt(0)
            p_l.paragraph_format.space_before = Pt(2)
            _styled_run(p_l, label, size=8, color=GREY)
            p_v = _next_paragraph()
            p_v.paragraph_format.space_after = Pt(0)
            _styled_run(p_v, value, bold=True, size=10, color=NAVY)

        # Spacer-Paragraph nach dem Hero/Sidebar-Block, damit Navy-Stripe
        # der Sidebar nicht direkt am folgenden Text klebt.
        spacer = self._doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(4)
        spacer.paragraph_format.space_before = Pt(0)
        _styled_run(spacer, " ", size=2)

        return table

    def lead(self, text: str):
        p = self._doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        _styled_run(p, text, bold=True, size=11, color=NAVY)
        return p

    def paragraph(self, text: str):
        p = self._doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _styled_run(p, text, size=10, color=TEXT)
        return p

    def labeled_paragraph(self, label: str, text: str):
        """Body-Absatz mit fettem Label-Präfix.

        Beispiel: labeled_paragraph('Die Aufgabe:', 'Im direkten Einfahrts...')
        rendert 'Die Aufgabe:' fett, gefolgt vom restlichen Text in normaler Stärke.
        """
        p = self._doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        _styled_run(p, label, bold=True, size=10, color=TEXT)
        _styled_run(p, " " + text, size=10, color=TEXT)
        return p

    def image_strip_grid(self,
                         rows: Sequence[Sequence[Tuple[str, str]]],
                         *, gap_cm: float = 0.25,
                         row_gap_pt: int = 4,
                         center_short_rows: bool = True,
                         image_width_cm: float | None = None,
                         space_before_pt: int = 4):
        """Mehrzeiliger Bildstreifen mit konsistenter Bildbreite über alle Reihen.

        Args:
            rows: Liste von Reihen; jede Reihe = Liste von (path, caption).
            gap_cm: Abstand zwischen Bildern in einer Reihe.
            row_gap_pt: Vertikaler Abstand zwischen den Reihen (Pt).
            center_short_rows: Wenn True, werden Reihen mit weniger Bildern
                als die längste Reihe horizontal zentriert.
            image_width_cm: Optional: Bildbreite vorgeben. Default = aus der
                längsten Reihe berechnet, sodass diese Reihe die Content-Breite
                füllt.
            space_before_pt: Vertikaler Abstand zur vorherigen Sektion
                (z. B. Drei-Akt-Story). Default 6 pt gibt eine spürbare
                Leerzeile, ohne die Seite zu sprengen.
        """
        if not rows:
            return None

        # Spacer-Paragraph vor dem Streifen, damit Text und Bilder nicht kleben.
        if space_before_pt > 0:
            spacer_top = self._doc.add_paragraph()
            spacer_top.paragraph_format.space_after = Pt(space_before_pt)
            spacer_top.paragraph_format.space_before = Pt(0)
            _styled_run(spacer_top, " ", size=2)

        max_cols = max(len(r) for r in rows)
        if image_width_cm is None:
            image_width_cm = round(
                (self.CONTENT_W - gap_cm * (max_cols - 1)) / max_cols, 2
            )

        for ri, row in enumerate(rows):
            n = len(row)
            used_w = n * image_width_cm + (n - 1) * gap_cm
            pad = round((self.CONTENT_W - used_w) / 2, 2) if (
                center_short_rows and used_w < self.CONTENT_W
            ) else 0

            # Spalten-Layout: [pad?] image gap image gap ... image [pad?]
            cols = []
            if pad > 0:
                cols.append(pad)
            for i in range(n):
                cols.append(image_width_cm)
                if i < n - 1:
                    cols.append(gap_cm)
            if pad > 0:
                cols.append(pad)

            table = self._doc.add_table(rows=1, cols=len(cols))
            _no_borders(table)
            _set_col_widths(table, cols)
            # Bild + Caption sollen nie über zwei Seiten gerissen werden.
            _cant_split_row(table.rows[0])

            offset = 1 if pad > 0 else 0
            for i, (path, caption) in enumerate(row):
                cell = table.cell(0, offset + i * 2)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                _cell_padding(cell, top=0, bottom=0, left=0, right=0)
                p_img = cell.paragraphs[0]
                p_img.paragraph_format.space_after = Pt(2)
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_img.add_run()
                _set_arial(run)
                run.add_picture(_orient(path), width=Cm(image_width_cm))
                p_cap = cell.add_paragraph()
                p_cap.paragraph_format.space_after = Pt(0)
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _styled_run(p_cap, caption, size=8, color=GREY)

            # Vertikaler Abstand zur nächsten Reihe via Spacer-Paragraph
            if ri < len(rows) - 1:
                spacer = self._doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(row_gap_pt)
                spacer.paragraph_format.space_before = Pt(0)
                _styled_run(spacer, " ", size=2)

        return None

    def image_strip(self, items: Sequence[Tuple[str, str]],
                    *, gap_cm: float = 0.25):
        """Drei Bilder mit Bildunterschrift, gleichmäßig verteilt."""
        n = len(items)
        if n == 0:
            return None
        col_w = round((self.CONTENT_W - gap_cm * (n - 1)) / n, 2)

        # Tabelle mit 2n-1 Spalten: image, gap, image, gap, image
        cols = []
        for i in range(n):
            cols.append(col_w)
            if i < n - 1:
                cols.append(gap_cm)
        table = self._doc.add_table(rows=1, cols=len(cols))
        _no_borders(table)
        _set_col_widths(table, cols)

        for i, (path, caption) in enumerate(items):
            cell = table.cell(0, i * 2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _cell_padding(cell, top=0, bottom=0, left=0, right=0)
            p_img = cell.paragraphs[0]
            p_img.paragraph_format.space_after = Pt(2)
            run = p_img.add_run()
            _set_arial(run)
            run.add_picture(_orient(path), width=Cm(col_w))
            p_cap = cell.add_paragraph()
            p_cap.paragraph_format.space_after = Pt(0)
            _styled_run(p_cap, caption, size=8, color=GREY)

        return table

    def footer(self, *, slug: str | None = None,
                url: str | None = None,
                cta_text: str | None = None,
                company_line: str = DEFAULT_COMPANY_LINE,
                base_url: str = DEFAULT_SLUG_BASE):
        """Word-Footer: Firmenzeile links + (CTA-Label +) QR-Code rechts.

        Genau einer von `slug` oder `url` muss gesetzt sein:
          - slug='brummer-lkw-zufahrt' → korodur.de/r/brummer-lkw-zufahrt
          - url='https://korodur.de/referenzen' → direkter Link

        `cta_text` ist optional und steht klein über dem QR-Code als
        Call-to-Action (z. B. „Mehr Referenzen entdecken").
        """
        if (slug is None) == (url is None):
            raise ValueError("footer() braucht genau eines: slug ODER url.")
        target_url = url if url else base_url.rstrip("/") + "/" + slug.lstrip("/")

        section = self._doc.sections[0]
        footer = section.footer

        # Bestehende Default-Footer-Paragraphen leeren
        for p in list(footer.paragraphs):
            p._p.getparent().remove(p._p)

        # QR-Bild rendern
        qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_M,
                           box_size=10, border=1)
        qr.add_data(target_url)
        qr.make(fit=True)
        img = qr.make_image(fill_color="#" + NAVY_HEX, back_color="white").convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)

        # Layout-Tabelle im Footer
        table = footer.add_table(rows=1, cols=2, width=Cm(self.CONTENT_W))
        _no_borders(table)
        _set_col_widths(table, [self.CONTENT_W - 3.0, 3.0])

        # Linke Zelle: Firmenzeile, vertikal unten
        c_info = table.cell(0, 0)
        c_info.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        _cell_padding(c_info, top=0, bottom=0, left=0, right=4)
        p_info = c_info.paragraphs[0]
        p_info.paragraph_format.space_after = Pt(0)
        _styled_run(p_info, company_line, size=8, color=DARKBLUE)

        # Rechte Zelle: optionaler CTA + QR (vertikal gestapelt, rechtsbündig)
        c_qr = table.cell(0, 1)
        c_qr.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        _cell_padding(c_qr, top=0, bottom=0, left=4, right=0)

        # Erster Paragraph: CTA-Text klein, rechtsbündig
        # space_after = Pt(6) gibt deutlich mehr Luft zum QR-Code darunter,
        # damit "Mehr Referenzen" nicht direkt am QR klebt.
        p_first = c_qr.paragraphs[0]
        p_first.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_first.paragraph_format.space_after = Pt(6)
        if cta_text:
            _styled_run(p_first, cta_text, italic=True, size=8, color=DARKBLUE)
            p_qr = c_qr.add_paragraph()
        else:
            p_qr = p_first
        p_qr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_qr.paragraph_format.space_after = Pt(0)
        run = p_qr.add_run()
        _set_arial(run)
        run.add_picture(buf, width=Mm(18))

    def save(self, path: str):
        os.makedirs(os.path.dirname(os.path.abspath(path)) or ".", exist_ok=True)
        self._doc.save(path)
        return path
